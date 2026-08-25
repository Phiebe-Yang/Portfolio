#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_horizontal_line(doc):
    """添加水平線"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def parse_markdown_to_docx(md_file, docx_file, image_file=None):
    """將 Markdown 文件轉換為 Word 文檔"""
    
    # 讀取 Markdown 文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 創建 Word 文檔
    doc = Document()
    
    # 設置文檔默認字體
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # 分行處理
    lines = content.split('\n')
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []
    image_inserted = False
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # 在附錄部分前插入圖片
        if not image_inserted and image_file and stripped.startswith('## 附錄'):
            p = doc.add_paragraph()
            p.add_run('系統架構流程圖\n').bold = True
            p_image = doc.add_paragraph()
            p_image.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p_image.add_run()
            run.add_picture(image_file, width=Inches(6.0))
            p_caption = doc.add_paragraph('RAG 系統完整架構和工作流程示意圖', style='List Bullet')
            p_caption.paragraph_format.left_indent = Inches(0.25)
            p_caption_run = p_caption.runs[0]
            p_caption_run.font.italic = True
            p_caption_run.font.size = Pt(9)
            doc.add_paragraph()
            image_inserted = True
        
        # 跳過空行
        if not stripped:
            if not in_code_block and not in_table:
                if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip():
                    doc.add_paragraph()
            continue
        
        # 代碼塊處理
        if stripped.startswith('```'):
            if in_code_block:
                # 結束代碼塊
                p = doc.add_paragraph()
                code_text = '\n'.join(code_block_lines)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                
                # 添加灰色背景效果
                p.paragraph_format.left_indent = Inches(0.25)
                
                in_code_block = False
                code_block_lines = []
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_lines.append(line.rstrip())
            continue
        
        # 標題
        if stripped.startswith('# ') and not stripped.startswith('##'):
            heading_text = stripped[2:]
            heading = doc.add_heading(heading_text, level=1)
        elif stripped.startswith('## ') and not stripped.startswith('###'):
            heading_text = stripped[3:]
            heading = doc.add_heading(heading_text, level=2)
        elif stripped.startswith('### ') and not stripped.startswith('####'):
            heading_text = stripped[4:]
            heading = doc.add_heading(heading_text, level=3)
        elif stripped.startswith('#### '):
            heading_text = stripped[5:]
            heading = doc.add_heading(heading_text, level=4)
        
        # 水平線
        elif stripped == '---' or stripped == '---':
            add_horizontal_line(doc)
        
        # 表格檢測
        elif '|' in stripped and not in_table:
            # 開始表格
            in_table = True
            table_lines = [stripped]
        
        elif '|' in stripped and in_table:
            table_lines.append(stripped)
        
        elif in_table and '|' not in stripped:
            # 結束表格 - 解析和添加
            if len(table_lines) >= 3:  # 最少需要標頭、分隔符、一行數據
                rows = []
                for table_line in table_lines:
                    cells = [cell.strip() for cell in table_line.split('|')]
                    cells = [cell for cell in cells if cell and cell != ':---' and not cell.replace('-', '').replace(':', '') == '']
                    if cells:
                        rows.append(cells)
                
                if rows:
                    # 創建表格（行數，列數）
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    # 填充表格數據
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
            
            in_table = False
            table_lines = []
            # 重新處理當前行（如果不是表格行）
            if '|' not in stripped:
                if stripped.startswith('- '):
                    doc.add_paragraph(stripped[2:], style='List Bullet')
                elif stripped.startswith('* '):
                    doc.add_paragraph(stripped[2:], style='List Bullet')
                elif re.match(r'^\d+\. ', stripped):
                    match = re.match(r'^(\d+)\. (.*)', stripped)
                    if match:
                        doc.add_paragraph(match.group(2), style='List Number')
                else:
                    doc.add_paragraph(stripped)
        
        # 列表項
        elif stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', stripped):
            match = re.match(r'^(\d+)\. (.*)', stripped)
            if match:
                doc.add_paragraph(match.group(2), style='List Number')
        
        # 普通段落（含粗體、斜體）
        else:
            p = doc.add_paragraph()
            
            # 用簡單的正則表達式解析粗體和斜體
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)', stripped)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('__') and part.endswith('__'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)
    
    # 保存文檔
    doc.save(docx_file)
    print('轉換成功!')
    print('檔案位置: ' + docx_file)

if __name__ == '__main__':
    md_path = r'C:\Users\Phiebe\OneDrive\Desktop\RAG\RAG建置說明書_正式版.md'
    docx_path = r'C:\Users\Phiebe\OneDrive\Desktop\RAG\RAG建置說明書_最終版.docx'
    image_path = r'C:\Users\Phiebe\OneDrive\Desktop\RAG\RAG_architecture_diagram.png'
    parse_markdown_to_docx(md_path, docx_path, image_path)
